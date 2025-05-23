import os
import re
import time
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from langchain.chat_models import ChatOpenAI
from openai import AsyncOpenAI

from ..models.report import ReportSection
from .base_agent import BaseAgent

WRITER_SYSTEM_PROMPT = """You are an expert content writer. Your task is to:
1. Write exceptionally comprehensive, detailed content DIRECTLY ABOUT THE USER'S REQUESTED TOPIC
2. Maximize the token count for each section without sacrificing quality
3. Provide in-depth explanations, examples, and analysis ABOUT THE TOPIC, not about what the sections are
4. Create a professional, authoritative tone throughout
5. Break down complex topics into clear subsections
6. Fully explore all facets of the topic with rich detail
7. Include comprehensive supporting evidence, statistics, and research findings
8. Use proper Markdown formatting for:
   - Headers (# for h1, ## for h2, etc.)
   - Emphasis (**bold** and *italic*)
   - Lists (- or 1. for ordered)
   - Tables (| header | header |)
   - Links [text](url)
   - Images ![caption](description) - REQUIRED: Include at least one image per section using this format

IMPORTANT CONTENT REQUIREMENTS:
1. FOCUS: NEVER explain what a section type (like "Executive Summary" or "Methodology") is supposed to be. Instead, ALWAYS write actual content directly addressing the user's requested topic.
   - INCORRECT: "An Executive Summary is a brief overview of a longer document..."
   - CORRECT: "This report examines [SPECIFIC TOPIC], finding that [KEY INSIGHTS ABOUT THE TOPIC]..."

2. LENGTH: Each section MUST be comprehensive and thorough, using at least 1000-1500 words per section
   - Provide extensive details, examples, case studies, and analysis ABOUT THE REQUESTED TOPIC
   - Break down topics into multiple well-developed paragraphs
   - Include nuanced analysis that explores multiple perspectives
   - Connect ideas across the document for a cohesive narrative

3. STRUCTURE: Each section should have a clear internal structure:
   - Begin with a comprehensive introduction to the section topic AS IT RELATES TO THE USER'S REQUESTED TOPIC
   - Develop 3-5 major points with substantial supporting content for each
   - Include relevant subsections with descriptive headers
   - End with a thorough conclusion that synthesizes key insights

4. DEPTH: Content must demonstrate expert-level understanding:
   - Include technical details appropriate to the audience
   - Provide historical context and future implications
   - Compare and contrast different approaches, methodologies, or viewpoints
   - Address potential criticisms or alternative perspectives

IMPORTANT IMAGE REQUIREMENTS:
1. Each section MUST include at least one relevant image using the syntax: ![caption](description)
   - This is a STRICT REQUIREMENT - responses without images will be rejected
   - Place the image markdown after the main content of each section
   - Images should be relevant to the section's topic and enhance understanding
   - You MUST include at least one image for EVERY section, no exceptions

2. Image descriptions must be detailed and specific, focusing on:
   - Professional visualizations (charts, graphs, diagrams)
   - Data-driven graphics (statistics, trends, comparisons)
   - Process flows (step-by-step illustrations)
   - Technical illustrations (system architectures, component interactions)
   - Conceptual diagrams (relationship maps, hierarchies)

3. Example image requests:
   ![AI Trading System Architecture](A professional technical diagram showing the components of an AI trading system, including data inputs, processing layers, and decision outputs, using a clean minimalist style with blue and gray color scheme)
   
   ![Risk Assessment Process](A detailed flowchart illustrating how AI analyzes financial risks, from data collection through risk scoring, with clear steps and decision points, using professional iconography)
   
   ![Customer Service Automation](An infographic showing how AI chatbots handle customer inquiries, with a timeline of interactions and key metrics, using a modern business visualization style)

4. Image Guidelines:
   - Keep descriptions between 50-100 words
   - Focus on clarity and professionalism
   - Avoid complex text in images
   - Use consistent visual style
   - Ensure business-appropriate content

CRITICAL: Your response MUST be extremely comprehensive, fully exploring each section topic in depth with maximum detail and information. You must maximize the token count while maintaining high-quality content. Each section should be thoroughly developed with multiple paragraphs, extensive analysis, and detailed explanations.

Remember: Your content will be rejected if it is not sufficiently detailed and comprehensive. MOST IMPORTANTLY, DO NOT EXPLAIN WHAT SECTIONS ARE SUPPOSED TO BE - WRITE ACTUAL CONTENT ABOUT THE REQUESTED TOPIC."""


class ContentWriterAgent(BaseAgent):
    """Agent responsible for generating report content."""

    def __init__(self, temperature: float = 0.3):
        """Initialize the content writer agent with the gpt-4o model and increased max tokens.

        Args:
            temperature (float): The temperature for model responses
        """
        super().__init__(model="gpt-4o", temperature=temperature)
        # Store temperature as instance variable
        self.temperature = temperature
        # Override the LLM with max_tokens parameter
        self.llm = ChatOpenAI(
            model="gpt-4o",
            temperature=temperature,
            api_key=os.getenv("OPENAI_API_KEY"),
            max_tokens=4096,  # Maximum allowed for gpt-4o
        )

    async def execute(self, task: Dict[str, Any]) -> str:
        """Execute the content writing task.
        
        Args:
            task (Dict[str, Any]): The content writing task
            
        Returns:
            str: The path to the generated document
        """
        # Extract task parameters
        structure = task["structure"]
        research = task["research"]
        title = structure.title  # Get title from structure instead of as a separate param
        include_images = task.get("include_images", True)  # Default to including images
        max_concurrent_tasks = task.get("max_concurrent_tasks", 10)  # Default concurrency limit
        
        # Extract the main topic from the research for consistent referencing
        main_topic = title  # Default to using the title as the main topic
        
        # Try to extract a more specific topic from research metadata if available
        for item in research:
            if isinstance(item, dict) and 'metadata' in item and 'question' in item['metadata']:
                main_topic = item['metadata']['question']
                self.logger.info(f"Extracted main topic from research: {main_topic}")
                break
                
        # Create images directory if including images
        images_dir = None
        if include_images:
            images_dir = "output/images"
            os.makedirs(images_dir, exist_ok=True)
            self.logger.info(f"Images enabled, using directory: {images_dir}")
        else:
            self.logger.info("Images disabled for this report")
        
        # Format title for filename
        filename = title.replace(" ", "_").replace(":", "_").replace("/", "_")
        output_path = f"output/{filename}.docx"
        
        # Create a new document
        doc = Document()
        
        # Add title
        doc.add_heading(title, 0)
        
        # Create a lock for document access
        import asyncio
        doc_lock = asyncio.Lock()
        
        # Save the initial document with just the title
        doc.save(output_path)
        self.logger.info(f"Initial document saved to {output_path}")
        
        # Process top-level sections in parallel with concurrency limit
        top_level_sections = structure.sections
        
        # Log the parallelization plan
        self.logger.info(f"Processing {len(top_level_sections)} top-level sections with max concurrency of {max_concurrent_tasks}")
        
        # Create tasks for each top-level section
        section_tasks = []
        for i, section in enumerate(top_level_sections):
            is_first_section = (i == 0)
            task = self._process_section_async(
                doc, 
                section, 
                research, 
                images_dir if include_images else None, 
                level=1, 
                is_first_section=is_first_section, 
                main_topic=main_topic,
                output_path=output_path,
                doc_lock=doc_lock
            )
            section_tasks.append(task)
        
        # Run tasks with concurrency limit
        await self._run_with_concurrency(section_tasks, max_concurrent_tasks)
        
        # Final save
        async with doc_lock:
            doc.save(output_path)
        
        self.logger.info(f"Document completed and saved to {output_path}")
        return output_path
        
    async def _run_with_concurrency(self, tasks, concurrency_limit):
        """Run tasks with a concurrency limit.
        
        Args:
            tasks: List of coroutines to run
            concurrency_limit: Maximum number of tasks to run concurrently
        """
        import asyncio
        semaphore = asyncio.Semaphore(concurrency_limit)
        
        async def _task_with_semaphore(task):
            async with semaphore:
                return await task
        
        return await asyncio.gather(*[_task_with_semaphore(task) for task in tasks])
        
    async def _process_section_async(self, doc, section, research, images_dir, level, is_first_section, main_topic, output_path, doc_lock):
        """Process a section asynchronously.
        
        Args:
            doc: The document to add the section to
            section: The section to process
            research: The research results
            images_dir: The directory to save images to
            level: The heading level
            is_first_section: Whether this is the first section
            main_topic: The main topic of the report
            output_path: The path to save the document to
            doc_lock: Lock for document access
        """
        self.logger.info(f"Starting processing of section: {section.title}")
        
        # Generate content if not already present
        if not section.content:
            start_time = time.time()
            self.logger.info(f"Generating content for section: {section.title}")
            content = await self._generate_content(section.title, research, include_images=(images_dir is not None), main_topic=main_topic)
            section.content = content
            elapsed = time.time() - start_time
            self.logger.info(f"Content generated for {section.title}, took {elapsed:.2f}s, length: {len(content)} characters")
        
        # Add section to document (protected by lock)
        async with doc_lock:
            # Add heading
            heading = doc.add_heading(section.title, level=level)
            
            # Add content (convert from markdown to docx)
            if section.content:
                await self._convert_markdown_to_docx(section.content, doc, images_dir)
            
            # Save progress after each section
            doc.save(output_path)
            self.logger.info(f"Progress saved after adding section: {section.title}")
        
        # Process subsections if any (in parallel)
        if section.subsections:
            subsection_tasks = []
            for i, subsection in enumerate(section.subsections):
                task = self._process_section_async(
                    doc,
                    subsection,
                    research,
                    images_dir,
                    level + 1,
                    is_first_section and i == 0,
                    main_topic,
                    output_path,
                    doc_lock
                )
                subsection_tasks.append(task)
            
            # Run subsection tasks with sensible concurrency - limit to 3 subsections at a time
            # to avoid overwhelming the system with too many nested tasks
            await self._run_with_concurrency(subsection_tasks, 3)
        
        self.logger.info(f"Completed processing of section: {section.title}")
        return section.title

    async def _convert_markdown_to_docx(
        self, markdown_text: str, doc: Document, images_dir: str
    ) -> None:
        """Convert markdown text to Word document format.

        Args:
            markdown_text (str): The markdown text to convert
            doc (Document): The Word document
            images_dir (str): Directory to save generated images
        """
        if not markdown_text:
            return

        self.logger.debug(
            f"Converting markdown to DOCX. Images directory: {images_dir}"
        )
        self.logger.debug(
            f"Markdown text sample (first 100 chars): {markdown_text[:100]}..."
        )

        # Check if markdown contains image syntax
        image_matches = re.findall(r"!\[(.+?)\]\((.+?)\)", markdown_text)
        self.logger.debug(f"Found {len(image_matches)} image references in markdown")
        for i, (caption, description) in enumerate(image_matches):
            self.logger.debug(
                f"Image {i+1}: Caption='{caption}', Description='{description}'"
            )

        # Pre-process to fix inconsistent numbered lists
        # Replace patterns like "41." with proper "1." formatting
        markdown_text = re.sub(r"^(\d+)\.\s", r"1. ", markdown_text, flags=re.MULTILINE)

        # Split into paragraphs
        paragraphs = markdown_text.strip().split("\n\n")
        self.logger.debug(f"Split markdown into {len(paragraphs)} paragraphs")

        for paragraph_index, paragraph_text in enumerate(paragraphs):
            # Skip empty paragraphs
            if not paragraph_text.strip():
                continue

            self.logger.debug(f"Processing paragraph {paragraph_index + 1}")

            # Headers
            header_match = re.match(r"^(#+)\s+(.+)$", paragraph_text)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2)
                p = doc.add_paragraph(text, style=f"Heading {level}")
                continue

            # Lists
            if paragraph_text.startswith(("- ", "* ")) or re.match(
                r"^\d+\.\s", paragraph_text
            ):
                lines = paragraph_text.split("\n")
                for line in lines:
                    if line.startswith(("- ", "* ")):
                        p = doc.add_paragraph(style="List Bullet")
                        text = line[2:]  # Remove list marker
                        # Process inline formatting within list items
                        self._apply_inline_formatting(p, text)
                    elif re.match(r"^\d+\.\s", line):
                        p = doc.add_paragraph(style="List Number")
                        # Extract text after the number and period
                        text = re.sub(r"^\d+\.\s", "", line)
                        # Process inline formatting within list items
                        self._apply_inline_formatting(p, text)
                continue

            # Tables
            if "|" in paragraph_text:
                lines = paragraph_text.strip().split("\n")
                if all("|" in line for line in lines):
                    # Extract table data
                    table_data = []
                    for line in lines:
                        if "---" in line:  # Skip separator line
                            continue
                        cells = [cell.strip() for cell in line.strip("|").split("|")]
                        table_data.append(cells)

                    # Create table
                    if table_data:
                        table = doc.add_table(
                            rows=len(table_data), cols=len(table_data[0])
                        )
                        table.style = "Table Grid"

                        # Fill table
                        for i, row in enumerate(table_data):
                            for j, cell in enumerate(row):
                                cell_paragraph = table.cell(i, j).paragraphs[0]
                                self._apply_inline_formatting(cell_paragraph, cell)

                        # Add spacing after table
                        doc.add_paragraph()
                        continue

            # Images
            image_match = re.search(r"!\[(.+?)\]\((.+?)\)", paragraph_text)
            if image_match and images_dir:
                caption, description = image_match.groups()
                self.logger.debug(
                    f"Processing image markdown - Caption: {caption}, Description: {description}"
                )

                # Generate and save image
                self.logger.debug(
                    f"Attempting to generate image for caption: {caption}"
                )
                image_path = await self._generate_and_save_image(description, caption)

                if image_path:
                    self.logger.debug(f"Image generated successfully at: {image_path}")
                    # Add image to document
                    self._add_image(
                        doc, {"path": image_path, "caption": caption, "size": "large"}
                    )
                    self.logger.debug(f"Added image to document: {image_path}")
                else:
                    # Add placeholder text if image generation failed
                    self.logger.error(f"Image generation failed for caption: {caption}")
                    p = doc.add_paragraph("[Image generation failed]")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Remove image markdown from text
                paragraph_text = paragraph_text.replace(image_match.group(0), "")

                # If no other content, continue to next paragraph
                if not paragraph_text.strip():
                    continue

            # Blockquotes
            if paragraph_text.startswith("> "):
                blockquote_text = paragraph_text.replace("> ", "")
                p = doc.add_paragraph(style="Quote")
                self._apply_inline_formatting(p, blockquote_text)
                continue

            # Code block (wrapped in ```code```)
            code_block_match = re.match(
                r"^```(.*?)\n(.*?)```$", paragraph_text, re.DOTALL
            )
            if code_block_match:
                self.logger.debug(f"Processing code block: {paragraph_text}")
                # Extract the code content
                language = code_block_match.group(1).strip()
                code_content = code_block_match.group(2)
                # Create a paragraph for the code block with monospace font
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)

                # Add the language as a label if provided
                if language:
                    language_run = p.add_run(f"{language}:\n")
                    language_run.bold = True
                    language_run.font.name = "Courier New"
                    language_run.font.size = Pt(9)

                # Add the code content
                code_run = p.add_run(code_content)
                code_run.font.name = "Courier New"
                code_run.font.size = Pt(9)
                continue

            # Regular paragraph
            p = doc.add_paragraph()
            self._apply_inline_formatting(p, paragraph_text)

    def _apply_inline_formatting(self, paragraph, text):
        """Apply inline formatting to a paragraph.

        Args:
            paragraph: The paragraph to add formatting to
            text: The text to format
        """
        # We'll use a recursive approach to handle nested formatting
        self._process_formatting(paragraph, text)

    def _process_formatting(self, paragraph, text):
        """Process text with potentially nested formatting.

        Args:
            paragraph: The paragraph to add formatting to
            text: The text to format
        """
        remaining_text = text

        while remaining_text:
            # Links pattern: [text](url)
            # Process links first to avoid conflicts with other formatting
            link_match = re.search(r"\[(.+?)\]\((.+?)\)", remaining_text)
            if link_match:
                # Add text before the link match
                before_text = remaining_text[: link_match.start()]
                if before_text:
                    self._process_formatting(paragraph, before_text)

                # Add the link text with formatting in the link text if present
                link_text, url = link_match.groups()
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True

                # Continue with text after the link match
                remaining_text = remaining_text[link_match.end() :]
                continue

            # Bold pattern: **text**
            bold_match = re.search(r"\*\*(.+?)\*\*", remaining_text)
            if bold_match:
                # Add text before the bold match
                before_text = remaining_text[: bold_match.start()]
                if before_text:
                    self._process_formatting(paragraph, before_text)

                # Add the bold text
                bold_content = bold_match.group(1)
                run = paragraph.add_run()
                run.bold = True

                # Process nested formatting within bold text
                inner_text = bold_content

                # Handle nested italic inside bold
                italic_match = re.search(
                    r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", inner_text
                )
                if italic_match:
                    # Add text before the italic match
                    before_italic = inner_text[: italic_match.start()]
                    if before_italic:
                        run.text += before_italic

                    # Add the italic text
                    italic_run = paragraph.add_run(italic_match.group(1))
                    italic_run.bold = True
                    italic_run.italic = True

                    # Add text after the italic match
                    after_italic = inner_text[italic_match.end() :]
                    if after_italic:
                        run.text += after_italic
                else:
                    # No nested formatting, add as is
                    run.text = inner_text

                # Continue with text after the bold match
                remaining_text = remaining_text[bold_match.end() :]
                continue

            # Italic pattern: *text*
            italic_match = re.search(
                r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", remaining_text
            )
            if italic_match:
                # Add text before the italic match
                before_text = remaining_text[: italic_match.start()]
                if before_text:
                    self._process_formatting(paragraph, before_text)

                # Add the italic text
                italic_content = italic_match.group(1)
                run = paragraph.add_run()
                run.italic = True

                # Process nested formatting within italic text
                inner_text = italic_content

                # Handle nested bold inside italic
                bold_match = re.search(r"\*\*(.+?)\*\*", inner_text)
                if bold_match:
                    # Add text before the bold match
                    before_bold = inner_text[: bold_match.start()]
                    if before_bold:
                        run.text += before_bold

                    # Add the bold text
                    bold_run = paragraph.add_run(bold_match.group(1))
                    bold_run.bold = True
                    bold_run.italic = True

                    # Add text after the bold match
                    after_bold = inner_text[bold_match.end() :]
                    if after_bold:
                        run.text += after_bold
                else:
                    # No nested formatting, add as is
                    run.text = inner_text

                # Continue with text after the italic match
                remaining_text = remaining_text[italic_match.end() :]
                continue

            # Code pattern: `text`
            code_match = re.search(r"`(.+?)`", remaining_text)
            if code_match:
                # Add text before the code match
                before_text = remaining_text[: code_match.start()]
                if before_text:
                    self._process_formatting(paragraph, before_text)

                # Add the code text
                run = paragraph.add_run(code_match.group(1))
                run.font.name = "Courier New"

                # Continue with text after the code match
                remaining_text = remaining_text[code_match.end() :]
                continue

            # If no matches found, add the remaining text as is
            paragraph.add_run(remaining_text)
            break

    async def _generate_and_save_image(
        self, description: str, caption: str
    ) -> Optional[str]:
        """Generate and save an image based on the description.

        Args:
            description (str): The description of the image to generate
            caption (str): Caption for the image

        Returns:
            Optional[str]: Path to the saved image, or None if generation failed
        """
        # Input validation
        if not description or len(description) < 10:
            self.logger.error("Image description is too short or empty")
            return None

        self.logger.debug(f"Generating image for caption: {caption}")
        self.logger.debug(f"Using description: {description}")

        try:
            # Use the ImageGenerationAgent to generate the image
            from .image_generation_agent import ImageGenerationAgent

            image_agent = ImageGenerationAgent()
            task = {
                "description": description,
                "caption": caption,
                # Use default settings for size, quality, and style
            }

            result = await image_agent.execute(task)

            if result["success"]:
                self.logger.debug(
                    f"Image generated successfully: {result['image_path']}"
                )
                return result["image_path"]
            else:
                self.logger.error(
                    f"Image generation failed: {result.get('error', 'Unknown error')}"
                )
                return None

        except Exception as e:
            self.logger.error(f"Error generating/saving image: {str(e)}")
            return None

    async def _add_list(self, doc: Document, items: List[str]) -> None:
        """Add a list to the document.

        Args:
            doc (Document): The Word document
            items (List[str]): The list items
        """
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            await self._convert_markdown_to_docx(item, doc, None)

    def _add_table(self, doc: Document, table_data: Dict[str, Any]) -> None:
        """Add a table to the document.

        Args:
            doc (Document): The Word document
            table_data (Dict[str, Any]): The table data
        """
        rows = table_data.get("data", [])
        if not rows:
            return

        # Create table
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"

        # Add data
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                table_cell = table.cell(i, j)
                # Create a paragraph in the cell and format its text
                p = table_cell.paragraphs[0]
                remaining_text = cell

                # Process inline formatting
                while remaining_text:
                    # Bold
                    bold_match = re.match(r"\*\*(.+?)\*\*(.*)$", remaining_text)
                    if bold_match:
                        run = p.add_run(bold_match.group(1))
                        run.bold = True
                        remaining_text = bold_match.group(2)
                        continue

                    # Italic
                    italic_match = re.match(r"\*(.+?)\*(.*)$", remaining_text)
                    if italic_match:
                        run = p.add_run(italic_match.group(1))
                        run.italic = True
                        remaining_text = italic_match.group(2)
                        continue

                    # Links
                    link_match = re.match(r"\[(.+?)\]\((.+?)\)(.*)$", remaining_text)
                    if link_match:
                        text, url, rest = link_match.groups()
                        run = p.add_run(text)
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        run.underline = True
                        remaining_text = rest
                        continue

                    # Regular text
                    regular_match = re.match(r"([^*\[]+)(.*)$", remaining_text)
                    if regular_match:
                        text, rest = regular_match.groups()
                        p.add_run(text)
                        remaining_text = rest
                        continue

                    # If no matches found, add remaining text and break
                    p.add_run(remaining_text)
                    break

        # Add caption if provided
        if caption := table_data.get("caption"):
            caption_para = doc.add_paragraph("Table: " + caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_image(self, doc: Document, image_data: Dict[str, Any]) -> None:
        """Add an image to the document.

        Args:
            doc (Document): The Word document
            image_data (Dict[str, Any]): The image data
        """
        if path := image_data.get("path"):
            try:
                self.logger.debug(f"Adding image to document: {path}")

                # Check if image file exists
                if not os.path.exists(path):
                    self.logger.error(f"Image file does not exist: {path}")
                    return

                # Check image file size
                file_size = os.path.getsize(path)
                self.logger.debug(f"Image file size: {file_size} bytes")

                # Add some spacing before image
                doc.add_paragraph()

                # Add image with proper sizing
                width = Inches(6)  # Default width
                if size := image_data.get("size"):
                    if size == "small":
                        width = Inches(3)
                    elif size == "large":
                        width = Inches(6.5)

                self.logger.debug(f"Adding image with width: {width}")
                picture = doc.add_picture(path, width=width)

                # Center the image
                paragraph = picture._parent
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.logger.debug(f"Image centered in document")

                # Add caption if provided
                if caption := image_data.get("caption"):
                    caption_para = doc.add_paragraph("Figure: " + caption)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_para.style = "Caption"
                    self.logger.debug(f"Added caption: {caption}")

                # Add some spacing after image
                doc.add_paragraph()
                self.logger.debug(f"Successfully added image to document: {path}")

            except Exception as e:
                self.logger.error(f"Error adding image {path}: {str(e)}")
                import traceback

                self.logger.error(traceback.format_exc())

    async def _generate_content(
        self,
        section_title: str,
        research: List[Dict[str, Any]],
        include_images: bool = True,
        main_topic: str = "",
    ) -> str:
        """Generate content for a section.

        Args:
            section_title (str): The title of the section
            research (List[Dict[str, Any]]): The research results
            include_images (bool): Whether to include images in the content
            main_topic (str): The main topic of the report

        Returns:
            str: The generated content
        """
        # Search for relevant research for this section
        section_research = []
        section_keywords = set(section_title.lower().split())

        for item in research:
            # Use get() with default value to handle missing 'title' key
            item_title = item.get("title", "")
            # If there's a 'section' key, use that as a fallback
            if not item_title and "section" in item:
                item_title = item.get("section", "")

            item_keywords = set(item_title.lower().split()) if item_title else set()
            # If there's any overlap in keywords, include this research
            if section_keywords & item_keywords or "all" in item_keywords:
                section_research.append(item)

        # If no specific research found, use all research
        if not section_research:
            section_research = research

        # Get target word count based on metadata if available
        target_word_count = "1000-1500"  # Default
        target_pages = 0

        for item in research:
            if isinstance(item, dict) and "metadata" in item:
                metadata = item.get("metadata", {})
                if "target_pages" in metadata:
                    target_pages = metadata.get("target_pages", 0)
                    break

        # Determine appropriate word count based on document total pages and section importance
        is_key_section = section_title.lower() in [
            "executive summary",
            "introduction",
            "findings",
            "conclusion",
            "recommendations",
        ]

        if target_pages > 0:
            # Approximate 500 words per page
            if target_pages <= 5:  # Short document
                target_word_count = "400-600" if is_key_section else "300-500"
            elif target_pages <= 10:  # Medium document
                target_word_count = "600-900" if is_key_section else "500-700"
            else:  # Long document
                target_word_count = "1000-1500" if is_key_section else "800-1200"

        image_instructions = (
            """
## IMAGE INSTRUCTIONS:
- REQUIRED: Include at least one image using markdown format ![caption](description)
- The image should be relevant to the section's topic and enhance understanding
- Image descriptions must be detailed and specific, focusing on:
  - Professional visualizations (charts, graphs, diagrams)
  - Data-driven graphics (statistics, trends, comparisons)
  - Process flows (step-by-step illustrations)
  - Technical illustrations (system architectures, component interactions)
  - Conceptual diagrams (relationship maps, hierarchies)
"""
            if include_images
            else ""
        )

        prompt = f"""
# Writing Task: Generate Comprehensive Content for "{section_title}" on the topic "{main_topic}"

## CRITICAL INSTRUCTION:
DO NOT explain what a "{section_title}" is supposed to be. Instead, write actual, substantive content about "{main_topic}" that belongs in this section.

## CONTENT REQUIREMENTS:
1. Create extremely detailed, in-depth content about "{main_topic}" for this section
2. Produce {target_word_count} words of high-quality, comprehensive content
3. Break down the topic into clear subsections with descriptive headers
4. Provide thorough analysis, examples, case studies, and evidence
{'' if not include_images else '5. Include at least one relevant image or diagram'}
6. Format properly using markdown

Your response MUST be formatted in well-structured Markdown, including:
- Clear **headings** (# for main headings) and **subheadings** (## or ###) to organize content
- **Bulleted lists** (- item) or **numbered lists** (1. item) for sequential information
- **Tables** using proper Markdown syntax (| Header | Header |) when presenting comparative data
- **Bold** (**text**) for emphasis on key points and terminology
- *Italic* (*text*) for definitions or secondary emphasis
- `Code blocks` for technical terms or snippets when relevant
- > Blockquotes for direct quotations from sources

## SECTION TOPIC:
{section_title} of {main_topic}

## RELEVANT RESEARCH:
{self._format_research_for_prompt(section_research)}

## SPECIAL INSTRUCTIONS:
- This is a professional document, so maintain appropriate tone and depth
- Target {target_word_count} words of valuable, high-quality content
{image_instructions}
- Break down complex topics into digestible parts while maintaining depth
- Connect this section to the overall document narrative
- Aim for comprehensive coverage that thoroughly explores all aspects of this topic

IMPORTANT: DO NOT write about what a "{section_title}" is or does in reports. Write ACTUAL CONTENT about "{main_topic}" appropriate for this section type.

Write exceptionally detailed content for this section now, maximizing thoroughness and information density:
"""

        response = await self._call_llm(WRITER_SYSTEM_PROMPT, prompt)

        # Check if there's an image in the content when images are required
        if include_images and "![" not in response:
            # If no image, add a warning and regenerate with stronger emphasis on images
            prompt_with_image_warning = (
                prompt
                + "\n\nWARNING: Your previous response did not include any images. YOU MUST INCLUDE AT LEAST ONE IMAGE using the format ![caption](description). This is a strict requirement."
            )
            response = await self._call_llm(
                WRITER_SYSTEM_PROMPT, prompt_with_image_warning
            )

        # Clean up response (remove markdown artifacts if any)
        response = re.sub(r"^```markdown", "", response)
        response = re.sub(r"```$", "", response)

        return response.strip()

    async def _call_llm(self, system_prompt: str, user_prompt: str) -> str:
        """Call the LLM with the given prompts.

        Args:
            system_prompt (str): The system prompt
            user_prompt (str): The user prompt

        Returns:
            str: The LLM response
        """
        self.logger.debug("Calling LLM with prompts:")
        self.logger.debug(f"System prompt: {system_prompt}")
        self.logger.debug(f"User prompt: {user_prompt}")

        try:
            # Initialize OpenAI client directly
            client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))

            # Make the API call directly
            response = await client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=self.temperature,
                max_tokens=4096,
            )

            return response.choices[0].message.content
        except Exception as e:
            self.logger.error(f"Error calling LLM: {str(e)}")
            raise

    def _format_research_for_prompt(self, research: List[Dict[str, Any]]) -> str:
        """Format research data for inclusion in a prompt.

        Args:
            research (List[Dict[str, Any]]): List of research items

        Returns:
            str: Formatted research string
        """
        if not research:
            return "No specific research available for this section."

        formatted_items = []
        for i, item in enumerate(research):
            # Extract key information
            title = item.get("title", "Untitled Research")
            content = item.get("content", "")
            source = item.get("source", "Unknown Source")

            # Format the research item
            formatted_item = f"""
RESEARCH ITEM #{i+1}:
TITLE: {title}
SOURCE: {source}
CONTENT: {content}
---"""
            formatted_items.append(formatted_item)

        return "\n".join(formatted_items)
