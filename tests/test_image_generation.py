import os
import pytest
import asyncio
import unittest.mock as mock
from dotenv import load_dotenv
from src.agents.content_writer_agent import ContentWriterAgent
from src.models.report import ReportStructure, ReportSection
from docx import Document

# Load environment variables from .env.test
load_dotenv(".env.test")

@pytest.fixture(autouse=True)
def setup_test_env():
    """Setup test environment."""
    # Create output directory
    os.makedirs("output/images", exist_ok=True)
    os.makedirs("output", exist_ok=True)

@pytest.fixture
def mock_image_generation():
    """Mock the image generation to return a test image path."""
    # Create a mock test image file
    test_image_path = "output/images/test_image.png"
    with open(test_image_path, "w") as f:
        f.write("test image content")
    
    # Create a mock for the _generate_and_save_image method
    with mock.patch.object(ContentWriterAgent, '_generate_and_save_image', autospec=True) as mock_gen:
        mock_gen.return_value = asyncio.Future()
        mock_gen.return_value.set_result(test_image_path)
        yield mock_gen

@pytest.fixture
def content_writer(mock_image_generation):
    """Fixture to create a ContentWriterAgent instance with mocked image generation."""
    return ContentWriterAgent()

@pytest.mark.asyncio
@pytest.mark.skip("Requires real OpenAI API key")
async def test_image_generation_basic():
    """Test basic image generation functionality."""
    agent = ContentWriterAgent()
    description = "A professional flowchart showing how AI analyzes financial data, with clear steps and decision points"
    caption = "AI Financial Analysis Process"
    
    # Generate and save image
    image_path = await agent._generate_and_save_image(description, caption)
    
    # Verify image was generated and saved
    assert image_path is not None
    assert os.path.exists(image_path)
    assert image_path.endswith(".png")
    assert "output/images" in image_path
    
    # Verify image file is not empty
    assert os.path.getsize(image_path) > 0

@pytest.mark.asyncio
@pytest.mark.skip("Requires real OpenAI API key")
async def test_image_generation_complex_description():
    """Test image generation with a more complex description."""
    agent = ContentWriterAgent()
    description = """A detailed technical diagram illustrating a modern AI trading system architecture. 
    Show multiple components including: data inputs (market feeds, news, social media), 
    processing layers (data cleaning, feature extraction, model inference), 
    and decision outputs (trade signals, risk metrics). 
    Use a clean, minimalist style with a blue and gray color scheme."""
    caption = "AI Trading System Architecture"
    
    image_path = await agent._generate_and_save_image(description, caption)
    
    assert image_path is not None
    assert os.path.exists(image_path)
    assert os.path.getsize(image_path) > 0

@pytest.mark.asyncio
async def test_image_generation_error_handling():
    """Test error handling in image generation."""
    agent = ContentWriterAgent()
    
    # Test with empty description
    empty_path = await agent._generate_and_save_image("", "Empty Test")
    assert empty_path is None
    
    # Test with very short description
    short_path = await agent._generate_and_save_image("test", "Short Test")
    assert short_path is None

@pytest.mark.asyncio
@pytest.mark.skip("Requires real OpenAI API key")
async def test_image_generation_concurrent():
    """Test concurrent image generation."""
    agent = ContentWriterAgent()
    descriptions = [
        ("A bar chart showing AI adoption rates across different financial sectors", "AI Adoption Rates"),
        ("A network diagram showing AI-powered fraud detection system", "Fraud Detection Network"),
        ("A timeline infographic of AI implementation stages in banking", "AI Implementation Timeline")
    ]
    
    # Generate images concurrently
    tasks = [agent._generate_and_save_image(desc, caption) for desc, caption in descriptions]
    results = await asyncio.gather(*tasks)
    
    # Verify all images were generated successfully
    for path in results:
        assert path is not None
        assert os.path.exists(path)
        assert os.path.getsize(path) > 0

@pytest.mark.asyncio
@pytest.mark.skip("Requires real OpenAI API key")
async def test_image_in_docx():
    """Test that generated images are properly added to the DOCX file."""
    agent = ContentWriterAgent()
    
    # Create a test document
    doc = Document()
    doc.add_heading("Test Document", 0)
    
    # Generate an image
    description = "A professional flowchart showing how AI analyzes financial data, with clear steps and decision points"
    caption = "AI Financial Analysis Process"
    
    # Generate and save image
    image_path = await agent._generate_and_save_image(description, caption)
    assert image_path is not None
    
    # Add image to document
    agent._add_image(doc, {"path": image_path, "caption": caption, "size": "large"})
    
    # Save document
    doc_path = "output/test_document.docx"
    doc.save(doc_path)
    
    # Verify document exists and has content
    assert os.path.exists(doc_path)
    assert os.path.getsize(doc_path) > 0
    
    # Open document and verify it has images
    doc = Document(doc_path)
    image_found = False
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_found = True
            break
    assert image_found, "No image found in the DOCX file"

@pytest.mark.asyncio
@pytest.mark.skip("Requires real OpenAI API key")
async def test_full_document_with_images():
    """Test generating a complete document with multiple images."""
    agent = ContentWriterAgent()
    
    # Create test structure with proper objects
    structure = ReportStructure(
        title="Test Report",
        sections=[
            ReportSection(
                title="Introduction",
                content="""# Introduction
This is a test section.

![AI Process Overview](A professional diagram showing AI data processing workflow with input, processing, and output stages)
"""
            ),
            ReportSection(
                title="Analysis",
                content="""# Analysis
This is another test section.

![Data Analysis](A detailed technical visualization showing data analysis pipeline with machine learning components)
"""
            )
        ],
        metadata={
            "template_type": "test",
            "total_sections": 2,
            "test_mode": True
        }
    )
    
    # Create test task
    task = {
        "structure": structure,
        "research": [],
        "include_images": True
    }
    
    # Generate document
    doc_path = await agent.execute(task)
    
    # Verify document exists
    assert os.path.exists(doc_path)
    assert os.path.getsize(doc_path) > 0
    
    # Open document and verify it has images
    doc = Document(doc_path)
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    
    # Should have at least 2 images (one per section)
    assert image_count >= 2, f"Expected at least 2 images, found {image_count}" 